#!/usr/bin/env python3
"""
md-to-docx.py — Convert a tailored Markdown resume to a formatted Word document.

Usage:
    python Resume/md-to-docx.py "Resume/Tailored/IEH resume-Feb2026-Exiger-CSL.md"

Output:
    Same path as input, with .docx extension.

Requires python-docx (installed in Resume/.venv):
    source Resume/.venv/bin/activate
    python md-to-docx.py <input.md>
"""

import sys
import copy
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── Paths ────────────────────────────────────────────────────────────────────

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(SCRIPT_DIR, "resume-template.docx")

# ── Colors ───────────────────────────────────────────────────────────────────

COLOR_TEXT  = RGBColor(0x33, 0x33, 0x33)
COLOR_EMAIL = RGBColor(0x40, 0x78, 0xC0)
COLOR_DIVIDER = "BFBFBF"

# ── Helpers ──────────────────────────────────────────────────────────────────

def clear_document(doc):
    """Remove all body paragraphs and tables from the document."""
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("p", "tbl", "sdt"):
            body.remove(child)


def set_paragraph_format(para, space_before=0, space_after=0,
                          left_indent=None, first_line_indent=None):
    fmt = para.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after  = Pt(space_after)
    if left_indent is not None:
        fmt.left_indent = Pt(left_indent)
    if first_line_indent is not None:
        fmt.first_line_indent = Pt(first_line_indent)


def add_run(para, text, font_name, font_size, bold=False, italic=False,
            color=None):
    run = para.add_run(text)
    run.font.name   = font_name
    run.font.size   = Pt(font_size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def apply_bullet_numbering(para, doc):
    """Apply the bullet numbering from the template document."""
    # Find the first numbering ID used in the template's styles
    # Fall back to a simple dash if none found
    numbering = doc.part.numbering_part
    if numbering is None:
        # No numbering available — use a dash prefix instead
        return False

    # Find abstractNum with a bullet format
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    root = numbering._element
    num_elements = root.findall("w:num", nsmap)
    if not num_elements:
        return False

    num_id = num_elements[0].get(qn("w:numId"))

    pPr = para._element.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), num_id)
    numPr.append(ilvl)
    numPr.append(numId_el)
    pPr.append(numPr)
    return True


def add_divider_border(cell):
    """Add a right border to a table cell to act as a vertical divider."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    right = OxmlElement("w:right")
    right.set(qn("w:val"),   "single")
    right.set(qn("w:sz"),    "4")
    right.set(qn("w:space"), "0")
    right.set(qn("w:color"), COLOR_DIVIDER)
    tcBorders.append(right)
    tcPr.append(tcBorders)


def remove_cell_borders(cell):
    """Remove all borders from a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tcBorders.append(el)
    tcPr.append(tcBorders)


# ── Section writers ───────────────────────────────────────────────────────────

def write_name(doc, text):
    para = doc.add_paragraph()
    set_paragraph_format(para, space_before=0, space_after=0)
    add_run(para, text.strip(), "Cambria", 26, bold=True, color=COLOR_TEXT)


def write_contact(doc, tokens):
    """tokens: list of contact strings; all rendered in COLOR_TEXT."""
    para = doc.add_paragraph()
    set_paragraph_format(para, space_before=0, space_after=6)
    for i, token in enumerate(tokens):
        if i > 0:
            add_run(para, "  |  ", "Helvetica Neue", 11, italic=True,
                    color=COLOR_TEXT)
        add_run(para, token.strip(), "Helvetica Neue", 11,
                italic=True, color=COLOR_TEXT)


def write_h2(doc, text):
    para = doc.add_paragraph()
    set_paragraph_format(para, space_before=6, space_after=4)
    add_run(para, text.strip(), "Helvetica Neue", 18, bold=True,
            color=COLOR_TEXT)


def write_h3(doc, text):
    para = doc.add_paragraph()
    set_paragraph_format(para, space_before=6, space_after=4)
    add_run(para, text.strip(), "Helvetica Neue", 14, bold=True, italic=True, color=COLOR_TEXT)


def write_h4(doc, text):
    """Bold-italic sub-heading for job title lines (e.g. **Expert Partner, Apr-2020**)."""
    para = doc.add_paragraph()
    set_paragraph_format(para, space_before=3, space_after=2)
    add_run(para, text.strip(), "Helvetica Neue", 11, bold=True, italic=True,
            color=COLOR_TEXT)


def write_bullet(doc, text, foundational=False):
    """Write a bullet paragraph with 0.2" hanging indent.

    The bullet is followed by a tab; a tab stop at 0.2" ensures the text
    starts — and wrapped lines return — to exactly 0.2" from the left margin.
    If foundational=True, the text before the first ': ' is rendered italic.
    """
    para = doc.add_paragraph()
    para.paragraph_format.space_before       = Pt(0)
    para.paragraph_format.space_after        = Pt(2)
    para.paragraph_format.left_indent        = Inches(0.2)
    para.paragraph_format.first_line_indent  = Inches(-0.2)

    # Tab stop at 0.2" (288 twips) so text aligns with wrapped lines
    pPr = para._element.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "left")
    tab.set(qn("w:pos"), "288")  # 0.2" = 288 twips
    tabs.append(tab)
    pPr.append(tabs)

    # Bullet character + tab (tab snaps to 0.2" stop)
    add_run(para, "\u2022\t", "Helvetica Neue", 11, color=COLOR_TEXT)

    if foundational and ": " in text:
        before, after = text.split(": ", 1)
        add_run(para, before + ": ", "Helvetica Neue", 11,
                italic=True, color=COLOR_TEXT)
        add_run(para, after, "Helvetica Neue", 11, color=COLOR_TEXT)
    else:
        add_run(para, text.strip(), "Helvetica Neue", 11, color=COLOR_TEXT)


def write_body_text(doc, text):
    """Plain paragraph (e.g. un-bulleted Profile text)."""
    para = doc.add_paragraph()
    set_paragraph_format(para, space_before=0, space_after=2)
    add_run(para, text.strip(), "Helvetica Neue", 11, color=COLOR_TEXT)


def write_skills_competencies_table(doc, competencies, skills):
    """Render Skills (left) and Competencies (right) in a 2-column table."""
    rows = max(len(competencies), len(skills))
    table = doc.add_table(rows=rows + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Remove all default table borders
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)

    # Header row: "Skills" (left) | "Competencies" (right)
    hdr_left  = table.rows[0].cells[0]
    hdr_right = table.rows[0].cells[1]

    for cell in (hdr_left, hdr_right):
        remove_cell_borders(cell)
    add_divider_border(hdr_left)

    for cell, label in ((hdr_left, "Skills"), (hdr_right, "Competencies")):
        para = cell.paragraphs[0]
        set_paragraph_format(para, space_before=6, space_after=4)
        add_run(para, label, "Helvetica Neue", 18, bold=True, color=COLOR_TEXT)

    # Content rows: skills on left, competencies on right
    for i in range(rows):
        row = table.rows[i + 1]
        left_cell  = row.cells[0]
        right_cell = row.cells[1]

        for cell in (left_cell, right_cell):
            remove_cell_borders(cell)
        add_divider_border(left_cell)

        def _write_table_bullet(cell, text):
            para = cell.paragraphs[0]
            para.paragraph_format.space_before      = Pt(0)
            para.paragraph_format.space_after       = Pt(2)
            para.paragraph_format.left_indent       = Inches(0.2)
            para.paragraph_format.first_line_indent = Inches(-0.2)
            pPr = para._element.get_or_add_pPr()
            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "left")
            tab.set(qn("w:pos"), "288")
            tabs.append(tab)
            pPr.append(tabs)
            add_run(para, "\u2022\t", "Helvetica Neue", 11, color=COLOR_TEXT)
            add_run(para, text, "Helvetica Neue", 11, color=COLOR_TEXT)

        if i < len(skills):
            _write_table_bullet(left_cell, skills[i])
        if i < len(competencies):
            _write_table_bullet(right_cell, competencies[i])


# ── Markdown parser ──────────────────────────────────────────────────────────

def parse_markdown(lines):
    """
    Parse resume Markdown into a list of (type, content) tuples.

    Types: name, contact, h2, h3, bullet, body, skills_competencies_table
    """
    elements = []
    i = 0
    n = len(lines)

    # Collect contact tokens (lines between H1 and first H2)
    contact_tokens = []
    in_contact = False

    # Track whether we're in the Foundational Roles subsection
    in_foundational = False

    # Buffer for Competencies/Skills table
    competencies = []
    skills       = []
    in_competencies = False
    in_skills       = False

    def flush_comp_skills():
        nonlocal competencies, skills, in_competencies, in_skills
        if competencies or skills:
            elements.append(("skills_competencies_table",
                             (list(competencies), list(skills))))
            competencies.clear()
            skills.clear()
        in_competencies = False
        in_skills = False

    while i < n:
        raw  = lines[i]
        line = raw.strip()
        i += 1

        if not line:
            continue

        # H1 — name
        if line.startswith("# ") and not line.startswith("## "):
            in_contact = True
            elements.append(("name", line[2:].strip()))
            continue

        # H2 — section label
        if line.startswith("## "):
            # Flush table buffer if we're leaving comp/skills sections
            label = line[3:].strip()
            if label.lower() not in ("competencies", "skills"):
                flush_comp_skills()
            if in_contact:
                if contact_tokens:
                    elements.append(("contact", contact_tokens[:]))
                    contact_tokens.clear()
                in_contact = False

            if label.lower() == "competencies":
                in_competencies = True
                in_skills = False
            elif label.lower() == "skills":
                in_skills = True
                in_competencies = False
            else:
                elements.append(("h2", label))
            continue

        # H3 — employer/role line
        if line.startswith("### "):
            flush_comp_skills()
            h3_text = line[4:].strip()
            in_foundational = h3_text.lower().startswith("foundational")
            elements.append(("h3", h3_text))
            continue

        # Bullet (* or -)
        if (line.startswith("* ") or line.startswith("- ")) and \
                not line.startswith("---"):
            text = line[2:].strip()
            # Skip placeholder bullets
            if text.startswith("_<") or text.startswith("**["):
                i_prev = i
                continue
            if in_competencies:
                competencies.append(text)
            elif in_skills:
                skills.append(text)
            elif in_contact:
                contact_tokens.append(text)
            else:
                elements.append(("bullet", text, in_foundational))
            continue

        # Contact lines (plain text between H1 and first H2)
        if in_contact:
            # Could be address, email, phone — treat each trailing-space
            # separated token as a contact token
            for token in line.split("  "):
                t = token.strip()
                if t:
                    contact_tokens.append(t)
            continue

        # HTML comments — skip
        if line.startswith("<!--"):
            continue

        # Horizontal rule — skip
        if line.startswith("---"):
            continue

        # Bold sub-heading (**text**)
        if line.startswith("**") and line.endswith("**") and len(line) > 4:
            elements.append(("h4", line[2:-2].strip()))
            continue

        # Plain body text
        if in_competencies or in_skills:
            continue  # non-bullet content inside these sections — skip
        elements.append(("body", line))

    flush_comp_skills()
    return elements


# ── Main ─────────────────────────────────────────────────────────────────────

def convert(md_path):
    # Derive output path
    base     = os.path.splitext(md_path)[0]
    out_path = base + ".docx"

    # Load template and clear its body
    doc = Document(TEMPLATE_PATH)
    clear_document(doc)

    # Set margins
    for section in doc.sections:
        section.top_margin    = Inches(0.80)
        section.bottom_margin = Inches(0.80)
        section.left_margin   = Inches(0.80)
        section.right_margin  = Inches(0.80)

    # Read and parse Markdown
    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    elements = parse_markdown(lines)

    # Write elements to document
    for element in elements:
        kind = element[0]
        if kind == "name":
            write_name(doc, element[1])
        elif kind == "contact":
            write_contact(doc, element[1])
        elif kind == "h2":
            write_h2(doc, element[1])
        elif kind == "h3":
            write_h3(doc, element[1])
        elif kind == "h4":
            write_h4(doc, element[1])
        elif kind == "bullet":
            write_bullet(doc, element[1], foundational=element[2])
        elif kind == "body":
            write_body_text(doc, element[1])
        elif kind == "skills_competencies_table":
            comp, skls = element[1]
            write_skills_competencies_table(doc, comp, skls)

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python md-to-docx.py <input.md>")
        sys.exit(1)
    convert(sys.argv[1])
